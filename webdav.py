import re
from webdav3.client import Client
from dotenv import load_dotenv
from os import getenv
from docx import Document
from datetime import datetime
import tempfile

load_dotenv()

WEBDAV_HOSTNAME = getenv('WEBDAV_HOSTNAME')
WEBDAV_LOGIN = getenv('WEBDAV_LOGIN')
WEBDAV_PASSWORD = getenv('WEBDAV_PASSWORD')

# Настройки подключения к WebDAV
options = {
    'webdav_hostname': WEBDAV_HOSTNAME,
    'webdav_login': WEBDAV_LOGIN,
    'webdav_password': WEBDAV_PASSWORD
}

MONTHS_RU = {
    '01_January': '01_ЯНВАРЬ',
    '02_February': '02_ФЕВРАЛЬ',
    '03_March': '03_МАРТ',
    '04_April': '04_АПРЕЛЬ',
    '05_May': '05_МАЙ',
    '06_June': '06_ИЮНЬ',
    '07_July': '07_ИЮЛЬ',
    '08_August': '08_АВГУСТ',
    '09_September': '09_СЕНТЯБРЬ',
    '10_October': '10_ОКТЯБРЬ',
    '11_November': '11_НОЯБРЬ',
    '12_December': '12_ДЕКАБРЬ'
}

client = Client(options)

def extract_date_time(message):
    match = re.search(r'(\d{4}-\d{2}-\d{2} \d{2}:\d{2}:\d{2})', message)
    if match:
        return datetime.strptime(match.group(1), '%Y-%m-%d %H:%M:%S')
    else:
        return datetime.now()

def generate_path(date_time):
    year = date_time.year
    month_eng = date_time.strftime('%m_%B')
    month_ru = MONTHS_RU.get(month_eng)
    file_name = date_time.strftime('%Y_%m_%d.docx')
    return f'/Дневник/{year}/{month_ru}/{file_name}'

def ensure_directories_exist(date_time):
    year = date_time.year
    month_name = date_time.strftime('%m_%B').upper()

    if not client.check(f'/Дневник/{year}'):
        client.mkdir(f'/Дневник/{year}')
    
    if not client.check(f'/Дневник/{year}/{month_name}'):
        client.mkdir(f'/Дневник/{year}/{month_name}')

def ensure_file_exists(date_time):
    path = generate_path(date_time)
    if not client.check(path):
        with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
            doc = Document()
            doc.save(tmp_file.name)
        client.upload_sync(remote_path=path, local_path=tmp_file.name)

def append_to_file(date_time, message_text):
    path = generate_path(date_time)

    ensure_directories_exist(date_time)
    ensure_file_exists(date_time)

    with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
        client.download_sync(remote_path=path, local_path=tmp_file.name)
        doc = Document(tmp_file.name)
    
    doc.add_paragraph(f'#Время {date_time.strftime("%H:%M:%S")}\n{message_text}')
   

    with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
        doc.save(tmp_file.name)
        client.upload_sync(remote_path=path, local_path=tmp_file.name)
